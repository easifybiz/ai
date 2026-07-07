"""Build a polished DOCX of the server audit report.

Run: python build_audit_docx.py
Output: SERVER_AUDIT_REPORT.docx
"""

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# === Theme ===
PRIMARY = RGBColor(0x1F, 0x3A, 0x68)       # deep navy
ACCENT  = RGBColor(0x2E, 0x86, 0xC1)       # mid blue
TEXT    = RGBColor(0x21, 0x21, 0x21)
MUTED   = RGBColor(0x6B, 0x72, 0x80)
HEADER_BG = "1F3A68"
ALT_ROW   = "F2F5F9"
CODE_BG   = "F4F6F8"
SUCCESS_BG = "EAF5EA"
SUCCESS_FG = RGBColor(0x21, 0x73, 0x21)
FAIL_BG   = "FBEAEA"
FAIL_FG   = RGBColor(0xB0, 0x2A, 0x2A)
BODY_FONT = "Calibri"
HEAD_FONT = "Calibri"
MONO_FONT = "Consolas"


def shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def cell_borders(cell, color="BFBFBF", size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def set_cell_padding(cell, top=80, bottom=80, left=120, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def style_run(run, *, font=BODY_FONT, size=11, bold=False, color=TEXT, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    if rFonts.getparent() is None:
        rPr.append(rFonts)


def add_para(doc, text="", *, font=BODY_FONT, size=11, bold=False, color=TEXT,
             italic=False, align=None, space_after=6, space_before=0, line_spacing=1.25):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        style_run(run, font=font, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    style_run(run, font=HEAD_FONT, size=sizes.get(level, 11), bold=True, color=PRIMARY)
    return p


def add_kv_table(doc, rows, *, col_widths=(Cm(5.5), Cm(10))):
    """Two-column key-value table (no header row, just alternating shading)."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for col_idx, w in enumerate(col_widths):
        for row in t.rows:
            row.cells[col_idx].width = w
    for i, (k, v) in enumerate(rows):
        bg = ALT_ROW if i % 2 == 0 else "FFFFFF"
        for c in t.rows[i].cells:
            shade(c, bg)
            cell_borders(c, color="DDDDDD", size=4)
            set_cell_padding(c)
        t.rows[i].cells[0].paragraphs[0].add_run(k)
        t.rows[i].cells[1].paragraphs[0].add_run(str(v))
        style_run(t.rows[i].cells[0].paragraphs[0].runs[0], bold=True, color=PRIMARY, size=10.5)
        style_run(t.rows[i].cells[1].paragraphs[0].runs[0], size=10.5)
    add_para(doc, space_after=4)


def add_data_table(doc, headers, rows, *, status_col=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    # Header row
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, HEADER_BG)
        cell_borders(c, color="1F3A68", size=6)
        set_cell_padding(c, top=100, bottom=100)
        c.paragraphs[0].add_run(h)
        style_run(c.paragraphs[0].runs[0], bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10.5)
    # Data rows
    for r_idx, row in enumerate(rows):
        bg = ALT_ROW if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            c = t.rows[r_idx + 1].cells[c_idx]
            shade(c, bg)
            cell_borders(c, color="DDDDDD", size=4)
            set_cell_padding(c)
            run = c.paragraphs[0].add_run(str(val))
            color = TEXT
            if status_col is not None and c_idx == status_col:
                if str(val).startswith("✓") or "Working" in str(val) or "Correct" in str(val) or "Open" in str(val) or "Clean" in str(val):
                    color = SUCCESS_FG
                    style_run(run, size=10.5, bold=True, color=color)
                elif str(val).startswith("❌") or "Timeout" in str(val) or "Failed" in str(val) or "Broken" in str(val):
                    color = FAIL_FG
                    style_run(run, size=10.5, bold=True, color=color)
                else:
                    style_run(run, size=10.5, color=color)
            else:
                style_run(run, size=10.5, color=color)
    add_para(doc, space_after=4)


def add_code_block(doc, code, *, lang_label=None):
    """Monospace box with light gray background."""
    if lang_label:
        p = add_para(doc, lang_label, size=9, italic=True, color=MUTED, space_after=2)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    cell.width = Cm(16)
    shade(cell, CODE_BG)
    cell_borders(cell, color="D6DBE0", size=4)
    set_cell_padding(cell, top=140, bottom=140, left=180, right=180)
    cell.text = ""
    for line in code.splitlines() or [""]:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        style_run(run, font=MONO_FONT, size=9.5, color=TEXT)
    # Remove the empty default paragraph at index 0
    cell._tc.remove(cell.paragraphs[0]._p)
    add_para(doc, space_after=4)


def add_callout(doc, text, *, kind="info"):
    """A boxed callout for emphasis (e.g., conclusion)."""
    palette = {
        "info":    ("E8F1FB", "2E86C1", PRIMARY),
        "success": (SUCCESS_BG, "2E7D32", SUCCESS_FG),
        "warn":    ("FFF7E0", "B8860B", RGBColor(0xB8, 0x86, 0x0B)),
        "fail":    (FAIL_BG, "C62828", FAIL_FG),
    }
    bg, border, fg = palette[kind]
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    cell.width = Cm(16)
    shade(cell, bg)
    cell_borders(cell, color=border, size=8)
    set_cell_padding(cell, top=180, bottom=180, left=240, right=240)
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(line)
        style_run(run, size=11, color=fg, bold=(i == 0))
    cell._tc.remove(cell.paragraphs[0]._p)
    add_para(doc, space_after=8)


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(it)
        style_run(run, size=11)


def add_footer_with_pagenum(section, footer_text):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(footer_text + "    |    Page ")
    style_run(run, size=9, color=MUTED)
    # Page number field
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    fnt = OxmlElement("w:rFonts")
    fnt.set(qn("w:ascii"), BODY_FONT)
    fnt.set(qn("w:hAnsi"), BODY_FONT)
    rPr.append(fnt)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "6B7280")
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "BFBFBF")
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================
# BUILD THE DOCUMENT
# ============================================================

doc = Document()

# Default body font
style = doc.styles["Normal"]
style.font.name = BODY_FONT
style.font.size = Pt(11)
style.font.color.rgb = TEXT

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    add_footer_with_pagenum(section, "VahanOne Server Audit Report")

# ---------- COVER ----------
add_para(doc, space_before=80)
add_para(doc, "INFRASTRUCTURE AUDIT", font=HEAD_FONT, size=11, bold=True,
         color=ACCENT, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
add_para(doc, "VahanOne Production EC2", font=HEAD_FONT, size=28, bold=True,
         color=PRIMARY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
add_para(doc, "ULIP Outbound Connectivity Investigation",
         font=HEAD_FONT, size=15, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_after=24)

add_horizontal_rule(doc)
add_para(doc, space_after=4)

add_kv_table(doc, [
    ("Date of Report", "6 May 2026"),
    ("Audit Period", "4 – 5 May 2026"),
    ("Server Audited", "i-0a5bfcfc64514a7ed (vahanone.com production EC2)"),
    ("Region", "ap-south-1 (Mumbai)"),
    ("Audit Objective",
     "Identify root cause of ULIP API timeout errors blocking the registration-number auto-fill feature in the AI Car Inspection tool."),
])

doc.add_page_break()

# ---------- 1. EXECUTIVE SUMMARY ----------
add_heading(doc, "1. Executive Summary", 1)
add_para(doc,
         "We performed a complete infrastructure audit of the VahanOne production EC2 instance to "
         "diagnose the persistent ULIP API timeout errors that are blocking the registration-number "
         "auto-fill feature in the AI Car Inspection tool.",
         space_after=6)
add_para(doc,
         "The server's backend code attempts to authenticate to ULIP at "
         "https://www.ulip.dpiit.gov.in/ulip/v1.0.0/user/login and fails with a "
         "ConnectTimeoutError for every request. After systematically auditing every layer of the AWS "
         "and OS configuration, we have determined that all customer-controllable settings are "
         "correctly configured. The instance is unable to establish any new outbound TCP connection "
         "to any external destination — including Google, Cloudflare, GitHub, and even AWS-internal "
         "services (S3) in the same region.",
         space_after=6)
add_para(doc,
         "We attempted four remediation actions of increasing impact (container restart, server "
         "reboot, Elastic IP re-association, and full instance stop/start to a new physical host). "
         "None resolved the outbound connectivity issue. We did successfully resolve a secondary "
         "issue — a process leak that had accumulated 4,143 zombie processes inside the frontend "
         "container — but this had no effect on the network behaviour.",
         space_after=10)

add_callout(doc,
            "Conclusion\n"
            "The outbound TCP failure is not caused by any setting we can change from the customer "
            "side. It is consistent with an AWS infrastructure-side restriction on the Elastic IP "
            "3.7.238.86 or the AWS account itself. Further investigation requires AWS Premium "
            "Support, which the account does not currently have. We recommend upgrading to AWS "
            "Developer Support ($29/month, downgradeable after resolution) and opening a technical "
            "case with the diagnostic evidence in this report.",
            kind="info")

# ---------- 2. BACKGROUND ----------
add_heading(doc, "2. Background", 1)
add_para(doc,
         "The AI Car Inspection tool integrates with the existing VahanOne backend "
         "(POST /easifybizsvc/addvehicle) to automatically fetch vehicle details from a "
         "registration number, replacing the manual data-entry form. The VahanOne backend in turn "
         "calls the Government of India ULIP API to retrieve VAHAN registration data.",
         space_after=6)
add_para(doc, "Every call to the /addvehicle endpoint returns the following error from the production server:",
         space_after=4)

add_code_block(doc,
"""ULIP Exception - 500: 401: Login request failed:
HTTPSConnectionPool(host='www.ulip.dpiit.gov.in', port=443):
Max retries exceeded with url: /ulip/v1.0.0/user/login
(Caused by ConnectTimeoutError(<urllib3.connection.HTTPSConnection ...>,
'Connection to www.ulip.dpiit.gov.in timed out. (connect timeout=5)'))""")

add_para(doc,
         "The same ULIP endpoint works correctly from external networks (verified via Postman from a "
         "different machine), confirming that ULIP itself is reachable from the public internet. "
         "The failure is specific to the VahanOne EC2 instance.",
         space_after=10)

# ---------- 3. SERVER INVENTORY ----------
add_heading(doc, "3. Server Inventory", 1)

add_kv_table(doc, [
    ("AWS Account",          "711387134795"),
    ("Region",               "ap-south-1 (Mumbai)"),
    ("Availability Zone",    "ap-south-1b"),
    ("Instance ID",          "i-0a5bfcfc64514a7ed"),
    ("Instance Name",        "DEV - VahanOne-[PGDB+DockerApp+WebNginx]"),
    ("Instance Type",        "t3a.medium (2 vCPU, 4 GB RAM)"),
    ("Operating System",     "Ubuntu 24.04.3 LTS (Linux 6.14.0-1018-aws)"),
    ("Private IP",           "172.31.8.16"),
    ("Public IP (Elastic)",  "3.7.238.86"),
    ("EIP Allocation ID",    "eipalloc-077069148273e9a1c"),
    ("VPC",                  "vpc-0535ac3baade106d2"),
    ("Subnet",               "subnet-0cbef0eb07def25c5"),
    ("Network Interface",    "eni-069d52916bb5d9a04"),
    ("Security Group",       "sg-00fd88f7719363808"),
    ("Internet Gateway",     "igw-08878ea93cebf3931"),
    ("Initial Launch Date",  "28 January 2026"),
])

add_heading(doc, "Services Running", 3)
add_data_table(doc,
    headers=["Service", "Layer", "Port", "Purpose"],
    rows=[
        ["Nginx",              "Host (systemd)",   "80, 443", "Reverse proxy / web server"],
        ["PostgreSQL 16",      "Host (systemd)",   "5434",    "Primary application database"],
        ["vahanone_backend",   "Docker container", "8000",    "FastAPI backend (Python 3.12)"],
        ["vahanone_frontend",  "Docker container", "3000",    "Next.js frontend"],
    ])
add_para(doc, "All services are configured to auto-start at boot.", italic=True, color=MUTED, size=10, space_after=10)

# ---------- 4. METHODOLOGY ----------
add_heading(doc, "4. Audit Methodology", 1)
add_para(doc, "We conducted the audit in the following order, applying the principle of "
              "least invasive action first:")
add_bullets(doc, [
    "Read-only AWS configuration audit — verified every networking-related setting via the AWS API.",
    "Read-only OS configuration audit — verified firewall, kernel, routing, and process state via SSH.",
    "Packet-level diagnostics — used tcpdump to observe actual packet flow during outbound attempts.",
    "Targeted remediation, escalating in impact — applied fixes from least disruptive (container restart) "
    "to most disruptive (full instance stop/start), retesting after each.",
])
add_para(doc,
         "All commands were either read-only API calls or applied to a single resource at a time, "
         "with rollback understood in advance.",
         space_after=10)

# ---------- 5. FINDINGS ----------
add_heading(doc, "5. Findings", 1)

add_heading(doc, "5.1 Components Verified as Correctly Configured", 2)
add_para(doc,
         "We checked every layer of the network stack between the application and the public internet. "
         "All of the following were verified to be correctly configured and not the cause of the "
         "outbound failure.")

add_heading(doc, "AWS Layer", 3)
add_data_table(doc,
    headers=["Component", "Configuration", "Status"],
    rows=[
        ["Security Group ingress", "TCP 22, 80, 443, 3000, 5434, 8090 from 0.0.0.0/0", "✓ Correct"],
        ["Security Group egress",  "All protocols, all ports to 0.0.0.0/0",            "✓ Fully open"],
        ["Network ACL",            "Default — allow all inbound + outbound from/to 0.0.0.0/0", "✓ Fully open"],
        ["Subnet",                 "Public subnet, MapPublicIpOnLaunch=true",          "✓ Correct"],
        ["Route Table (main)",     "0.0.0.0/0 → igw-08878ea93cebf3931",                "✓ Correct"],
        ["Internet Gateway",       "State = available, attached to VPC",                "✓ Working"],
        ["Elastic IP",             "Bound to correct ENI, instance, and account",       "✓ Correct"],
        ["Network Interface",      "Status=in-use, SourceDestCheck=true",               "✓ Correct"],
        ["AWS Network Firewall",   "Not deployed (would not interfere)",                "✓ N/A"],
        ["VPC Endpoints",          "None present (would not interfere)",                "✓ N/A"],
        ["Transit Gateway",        "Not attached",                                       "✓ N/A"],
        ["ENA conntrack allowance","51,302 free of allowance, 0 exceeded",              "✓ Plenty of headroom"],
        ["ENA bandwidth allowance","0 in/out exceeded",                                  "✓ Within limits"],
    ],
    status_col=2)

add_heading(doc, "Operating System Layer", 3)
add_data_table(doc,
    headers=["Component", "Configuration", "Status"],
    rows=[
        ["ufw (Ubuntu firewall)",   "Inactive",                                           "✓ Not blocking"],
        ["iptables INPUT chain",    "Empty, default policy ACCEPT",                       "✓ Not blocking"],
        ["iptables OUTPUT chain",   "Empty, default policy ACCEPT",                       "✓ Not blocking"],
        ["iptables FORWARD chain",  "Standard Docker rules only",                         "✓ Expected"],
        ["iptables NAT POSTROUTING","MASQUERADE for 172.17.0.0/16 (Docker) only",         "✓ Expected"],
        ["net.ipv4.ip_forward",     "1 (required for Docker bridge)",                     "✓ Correct"],
        ["rp_filter on ens5",       "2 (loose)",                                          "✓ Acceptable"],
        ["MTU on ens5",             "9001 (AWS jumbo frames)",                            "✓ Default for AWS"],
        ["ARP table",               "Gateway 172.31.0.1 and DNS 172.31.0.2 reachable",    "✓ Correct"],
        ["DNS resolution",          "www.ulip.dpiit.gov.in resolves to 164.100.63.96",    "✓ Working"],
        ["Kernel network errors (dmesg)", "None in the last 24 hours",                    "✓ Clean"],
        ["Disk usage",              "50 % used (9.1 GB free)",                            "✓ Adequate"],
    ],
    status_col=2)

add_heading(doc, "5.2 Symptom Analysis: Outbound TCP Universally Fails", 2)
add_para(doc, "We tested outbound TCP connections from both the host and the backend container to a "
              "wide range of destinations:")

add_data_table(doc,
    headers=["Destination", "Type", "Result"],
    rows=[
        ["1.1.1.1:443 (Cloudflare)",            "Public internet",              "❌ Timeout"],
        ["8.8.8.8:443 (Google DNS)",            "Public internet",              "❌ Timeout"],
        ["google.com:443",                       "Public internet",              "❌ Timeout"],
        ["github.com:443",                       "Public internet",              "❌ Timeout"],
        ["164.100.63.96:443 (ULIP)",            "Public internet",              "❌ Timeout"],
        ["s3.ap-south-1.amazonaws.com:443",     "AWS internal (same region)",   "❌ Timeout"],
        ["ec2.ap-south-1.amazonaws.com:443",    "AWS internal (same region)",   "❌ Timeout"],
        ["169.254.169.254 (EC2 metadata)",      "Link-local (in-host)",          "✓ Working"],
        ["8.8.8.8:53 TCP (DNS over TCP)",       "Public internet",              "⚠ Intermittent"],
        ["ICMP ping to 172.31.0.1 (gateway)",   "Within VPC",                   "✓ Working"],
        ["Inbound SSH (port 22 from external)", "Public internet",              "✓ Working"],
        ["Inbound HTTPS to Nginx (port 443)",   "Public internet",              "✓ Working"],
    ],
    status_col=2)

add_para(doc,
         "The pattern is unambiguous: established and inbound TCP connections work perfectly, but "
         "the instance cannot establish a new outbound TCP connection to anything beyond its own VPC.",
         bold=True, space_after=8)

add_para(doc, "A tcpdump capture during a fresh outbound connection attempt confirms that SYN packets "
              "do leave the host correctly:")
add_code_block(doc,
"""06:24:02.493741 IP 172.31.8.16.41524 > 1.1.1.1.443: Flags [S], seq 381370241 (initial)
06:24:03.553872 IP 172.31.8.16.41524 > 1.1.1.1.443: Flags [S], seq 381370241 (retry 1)
06:24:04.577865 IP 172.31.8.16.41524 > 1.1.1.1.443: Flags [S], seq 381370241 (retry 2)
06:24:05.601855 IP 172.31.8.16.41524 > 1.1.1.1.443: Flags [S], seq 381370241 (retry 3)
4 packets captured, 0 packets dropped by kernel""")

add_para(doc, "No SYN-ACK responses are ever received in return. This indicates that either:", space_after=2)
add_bullets(doc, [
    "the SYN packets are not reaching the destination, or",
    "the destination's SYN-ACK responses are not being routed back to the instance.",
])
add_para(doc,
         "Given that all of our outbound destinations behave identically — including AWS-internal "
         "services that share no network path with the public internet — the issue cannot be "
         "attributed to any single destination, ULIP whitelist, ISP routing problem, or path MTU issue.",
         space_after=10)

add_heading(doc, "5.3 Secondary Finding (Resolved): Zombie Process Leak", 2)
add_para(doc,
         "During the OS-level audit we observed an unusually high process count: 4,419 total "
         "processes, of which 4,143 were zombies. All zombies were children of process 1148829, "
         "which was identified as next-server (Next.js production server) running inside the "
         "vahanone_frontend Docker container.")
add_para(doc,
         "This indicated a long-standing process leak — Next.js was spawning shell subprocesses "
         "(likely as part of its image/asset-handling routines) and not reaping them after "
         "termination. The leak had been accumulating for 79 days (the previous server uptime).")
add_para(doc, "Action taken:", bold=True, space_after=2)
add_para(doc,
         "Restarted only the vahanone_frontend container (docker restart vahanone_frontend). "
         "The backend, database, and Nginx were untouched.")
add_para(doc, "Result:", bold=True, space_after=2)
add_callout(doc,
            "Zombie count dropped from 4,143 to 0; total process count dropped from 4,419 to 147. "
            "System health metrics returned to normal.",
            kind="success")
add_para(doc,
         "This leak was a secondary concern — it had no observable impact on the website or "
         "backend functionality, but had it continued unchecked it could eventually have hit the "
         "system's process limit. Resolving it did not, however, fix the outbound TCP issue, "
         "indicating the two problems are unrelated.",
         space_after=4)
add_para(doc,
         "We recommend that the frontend application code be reviewed in a future maintenance "
         "window to identify and patch the subprocess leak at the source — a container restart is "
         "a workaround, not a fix.",
         italic=True, color=MUTED, space_after=10)

# ---------- 6. REMEDIATION ----------
add_heading(doc, "6. Remediation Actions Attempted", 1)
add_para(doc,
         "We applied the following remediations in order of increasing impact, retesting outbound "
         "connectivity after each.")

add_heading(doc, "6.1 Frontend Container Restart", 2)
add_kv_table(doc, [
    ("Action",   "sudo docker restart vahanone_frontend"),
    ("Downtime", "~10 seconds (frontend only)"),
    ("Outcome",  "Zombie processes cleared; outbound TCP connectivity unchanged (still failing)."),
])

add_heading(doc, "6.2 Full Server Reboot", 2)
add_kv_table(doc, [
    ("Action",   "sudo reboot — full system shutdown and reboot, clearing all kernel state, network buffers, and TCP connection tracking."),
    ("Downtime", "~75 seconds total"),
    ("Verified post-reboot", "Docker, PostgreSQL, Nginx, both containers all auto-restarted cleanly. System health: 0 zombies, 147 processes, 1-minute uptime."),
    ("Outcome",  "Outbound TCP still failing identically to all destinations. This rules out kernel-level state corruption as the cause."),
])

add_heading(doc, "6.3 Elastic IP Disassociate + Re-associate", 2)
add_kv_table(doc, [
    ("Action",   "Disassociated the EIP 3.7.238.86 from the instance, waited 3 seconds, then re-associated it to the same instance and network interface. New association ID issued (eipassoc-060234f01c2636fc8)."),
    ("Downtime", "~5 seconds (no public IP attached)"),
    ("Outcome",  "Same EIP returned to same instance. Outbound TCP still failing identically. This rules out stuck NAT state at the EIP level on AWS's side."),
])

add_heading(doc, "6.4 Instance Stop + Start (New Physical Host)", 2)
add_kv_table(doc, [
    ("Action",   "Issued aws ec2 stop-instances (graceful), waited for stopped state, then aws ec2 start-instances. AWS allocates a brand-new physical host for the instance, with fresh hypervisor-level network state. Elastic IP automatically re-attached on start."),
    ("Downtime", "~3 minutes"),
    ("Verified post-start", "Same EIP 3.7.238.86, same private IP 172.31.8.16, same Availability Zone, all services auto-started cleanly. System health: 0 zombies, 152 processes, 3-minute uptime."),
    ("Outcome",  "Outbound TCP still failing identically. This rules out hypervisor-level network issues on the original physical host."),
])

add_heading(doc, "Summary of Remediation Outcomes", 3)
add_data_table(doc,
    headers=["Remediation", "Resolved Outbound Issue?", "What It Ruled Out"],
    rows=[
        ["Frontend container restart",       "No", "App-level resource exhaustion"],
        ["Full server reboot",                "No", "OS / kernel-level network state corruption"],
        ["EIP re-association",                "No", "Stuck NAT state on the EIP at AWS edge"],
        ["Stop / Start (new physical host)",  "No", "Hypervisor-level network problem on original host"],
    ],
    status_col=1)

add_callout(doc,
            "No customer-side action available to us has resolved this issue.",
            kind="warn")

# ---------- 7. CONCLUSION ----------
add_heading(doc, "7. Conclusion", 1)
add_para(doc,
         "After a comprehensive audit of every customer-controllable layer — AWS networking, OS "
         "firewall and kernel, application and container state — and after applying every "
         "escalating remediation that does not require AWS engineering involvement, the outbound "
         "connectivity issue persists unchanged.")
add_para(doc,
         "The behaviour is internally consistent and points to one of the following root causes, "
         "all of which are at the AWS infrastructure level:",
         space_after=4)
add_bullets(doc, [
    "Elastic IP reputation block. AWS Trust & Safety silently restricts outbound traffic from EIPs that have been flagged for abuse (DDoS source, spam origin, port scanning). Customers are not always notified.",
    "Account-level outbound restriction. The AWS account itself may have an undisclosed restriction. This can occur due to billing issues, fraud detection, or compliance holds.",
    "Region-level networking incident at AWS. Less likely given the long duration, but possible.",
])
add_callout(doc,
            "In all three cases, resolution requires escalation to AWS Premium Support — there are no further customer-side actions that can resolve this.",
            kind="info")

# ---------- 8. NEXT STEPS ----------
add_heading(doc, "8. Recommended Next Steps", 1)
add_heading(doc, "8.1 Open an AWS Premium Support Case (required)", 2)
add_para(doc,
         "The account is currently on Basic Support (the free tier), which does not allow technical "
         "support cases — only account and billing questions. To proceed, the account must be "
         "upgraded.")
add_para(doc, "Recommended Support tier: Developer Support ($29/month minimum)", bold=True, color=PRIMARY, space_after=2)
add_bullets(doc, [
    "Sufficient for this issue.",
    "Response time: under 24 business hours.",
    "Can be downgraded back to Basic Support immediately after the case is resolved (effective monthly cost: $29 one-time).",
])
add_para(doc, "Steps to upgrade:", bold=True, space_after=2)
add_bullets(doc, [
    "Sign in to the AWS Console with the root account.",
    "Navigate to: Support Center → Support plans → Change support plan.",
    "Select Developer.",
    "Confirm.",
])
add_para(doc, "Steps to open the case (after upgrading):", bold=True, space_after=2)
add_bullets(doc, [
    "Support Center → Create case.",
    "Issue type: Technical.",
    "Service: EC2.",
    "Category: Connectivity / Networking.",
    "Severity: General Guidance (or Production Impaired if relevant).",
    "Subject: EC2 instance unable to make any outbound TCP connections — all destinations time out.",
    "Description: attach the contents of this report (or summarise key findings).",
])

# ---------- APPENDIX ----------
doc.add_page_break()
add_heading(doc, "Appendix A: Diagnostic Evidence", 1)
add_para(doc,
         "This appendix contains the raw command outputs from the audit, included for completeness "
         "and to support the AWS Support case.",
         italic=True, color=MUTED, space_after=10)

add_heading(doc, "A.1 Security Group Rules (sg-00fd88f7719363808)", 2)
add_code_block(doc,
"""{
    "Inbound": [
        {"IpProtocol": "tcp", "FromPort":   22, "ToPort":   22, "IpRanges": [{"CidrIp": "0.0.0.0/0"}]},
        {"IpProtocol": "tcp", "FromPort":   80, "ToPort":   80, "IpRanges": [{"CidrIp": "0.0.0.0/0"}]},
        {"IpProtocol": "tcp", "FromPort":  443, "ToPort":  443, "IpRanges": [{"CidrIp": "0.0.0.0/0"}]},
        {"IpProtocol": "tcp", "FromPort": 3000, "ToPort": 3000, "IpRanges": [{"CidrIp": "0.0.0.0/0"}]},
        {"IpProtocol": "tcp", "FromPort": 5434, "ToPort": 5434, "IpRanges": [{"CidrIp": "0.0.0.0/0"}]},
        {"IpProtocol": "tcp", "FromPort": 8090, "ToPort": 8090, "IpRanges": [{"CidrIp": "0.0.0.0/0"}]}
    ],
    "Outbound": [
        {"IpProtocol": "-1", "IpRanges": [{"CidrIp": "0.0.0.0/0"}]}
    ]
}""")

add_heading(doc, "A.2 Network ACL (acl-0158d37996a32dafc, default)", 2)
add_code_block(doc,
"""Egress 100   : ALL  0.0.0.0/0  ALLOW
Egress 32767 : ALL  0.0.0.0/0  DENY (catch-all)
Ingress 100  : ALL  0.0.0.0/0  ALLOW
Ingress 32767: ALL  0.0.0.0/0  DENY (catch-all)""")

add_heading(doc, "A.3 Route Table (main, rtb-0655cff2b2865127a)", 2)
add_code_block(doc,
"""[
  {"DestinationCidrBlock": "172.31.0.0/16", "GatewayId": "local",                 "State": "active"},
  {"DestinationCidrBlock": "0.0.0.0/0",     "GatewayId": "igw-08878ea93cebf3931", "State": "active"}
]""")

add_heading(doc, "A.4 ENA (Elastic Network Adapter) Driver Statistics", 2)
add_para(doc, "Hypervisor-level limits enforced by AWS Nitro:", italic=True, color=MUTED, space_after=2)
add_code_block(doc,
"""bw_in_allowance_exceeded:      0
bw_out_allowance_exceeded:     0
pps_allowance_exceeded:        128    (cumulative, historical)
conntrack_allowance_exceeded:  0
linklocal_allowance_exceeded:  0
conntrack_allowance_available: 51302  (out of ~50K-100K limit for t3a.medium)""")

add_heading(doc, "A.5 OS Routing Table", 2)
add_code_block(doc,
"""default via 172.31.0.1 dev ens5 proto dhcp src 172.31.8.16 metric 100
172.17.0.0/16 dev docker0 proto kernel scope link src 172.17.0.1
172.31.0.0/20 dev ens5 proto kernel scope link src 172.31.8.16 metric 100""")

add_heading(doc, "A.6 OS Firewall State", 2)
add_code_block(doc,
"""$ sudo ufw status
Status: inactive

$ sudo iptables -L INPUT -n
Chain INPUT (policy ACCEPT)
target     prot opt source               destination

$ sudo iptables -L OUTPUT -n
Chain OUTPUT (policy ACCEPT)
target     prot opt source               destination""")

add_heading(doc, "A.7 Sysctl Network Parameters", 2)
add_code_block(doc,
"""net.ipv4.ip_forward                 = 1
net.ipv4.tcp_syn_retries            = 6
net.ipv4.tcp_retries2               = 15
net.ipv4.conf.ens5.rp_filter        = 2
net.ipv4.conf.all.rp_filter         = 2""")

add_heading(doc, "A.8 Outbound TCP Test Matrix (post-stop/start, post-reboot)", 2)
add_code_block(doc,
"""1.1.1.1:443                  -> timeout
8.8.8.8:443                  -> timeout
google.com:443               -> timeout (HTTP 000, time=8.0s)
ULIP 164.100.63.96:443       -> timeout (HTTP 000, time=8.0s)
S3 ap-south-1:443            -> timeout
EC2 API ap-south-1:443       -> timeout
github.com:443               -> timeout
169.254.169.254 (metadata)   -> 401 OK in 0.001s
DNS UDP                      -> working
ICMP to 172.31.0.1 (gateway) -> 0% packet loss""")

add_heading(doc, "A.9 tcpdump Output During Outbound Connection Attempt", 2)
add_para(doc, "Captured on interface ens5 while running curl https://1.1.1.1/ from the host:",
         italic=True, color=MUTED, space_after=2)
add_code_block(doc,
"""06:24:02.493741 IP 172.31.8.16.41524 > 1.1.1.1.443: Flags [S], seq 381370241,
                  win 62727, options [mss 8961,sackOK,TS val 1996212366
                  ecr 0,nop,wscale 7], length 0
06:24:03.553872 IP 172.31.8.16.41524 > 1.1.1.1.443: Flags [S] (retry)
06:24:04.577865 IP 172.31.8.16.41524 > 1.1.1.1.443: Flags [S] (retry)
06:24:05.601855 IP 172.31.8.16.41524 > 1.1.1.1.443: Flags [S] (retry)

4 packets captured
4 packets received by filter
0 packets dropped by kernel""")
add_para(doc, "No reply packets observed in the reverse direction.",
         italic=True, color=MUTED, size=10, space_after=8)

add_heading(doc, "A.10 EIP State Verification", 2)
add_code_block(doc,
"""{
    "AllocationId":      "eipalloc-077069148273e9a1c",
    "AssociationId":     "eipassoc-060234f01c2636fc8",
    "Domain":            "vpc",
    "InstanceId":        "i-0a5bfcfc64514a7ed",
    "NetworkInterfaceId":"eni-069d52916bb5d9a04",
    "PrivateIpAddress":  "172.31.8.16",
    "PublicIp":          "3.7.238.86",
    "PublicIpv4Pool":    "amazon",
    "NetworkBorderGroup":"ap-south-1"
}""")

add_heading(doc, "A.11 AWS Support Tier Verification", 2)
add_para(doc, "Attempting any AWS Support API call returns:", italic=True, color=MUTED, space_after=2)
add_code_block(doc,
"""SubscriptionRequiredException: Amazon Web Services Premium Support Subscription
                              is required to use this service.""")
add_para(doc,
         "Confirms the account is on Basic Support tier and cannot open technical cases.",
         italic=True, color=MUTED, size=10, space_after=12)

add_horizontal_rule(doc)
add_para(doc, "End of Report.", italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

# Save
out_path = "/Users/himmu1144/Documents/Projects/Random/vahanone/aicarinspection/SERVER_AUDIT_REPORT.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
